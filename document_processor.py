import os
import hashlib
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pandas as pd
from openpyxl import load_workbook
import re
from typing import List, Dict, Any, Tuple

class DocumentProcessor:
    """Handles parsing and text extraction from various document formats"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.xlsx'}
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if the file format is supported"""
        _, ext = os.path.splitext(filename.lower())
        return ext in self.supported_formats
    
    def calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of the file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process a document and extract text content"""
        _, ext = os.path.splitext(filename.lower())
        
        try:
            if ext == '.pdf':
                return self._process_pdf(file_path)
            elif ext == '.docx':
                return self._process_docx(file_path)
            elif ext == '.xlsx':
                return self._process_xlsx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'extracted_text': '',
                'chunks': [],
                'metadata': {}
            }
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF"""
        doc = fitz.open(file_path)
        extracted_text = ""
        chunks = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text()
            
            if page_text.strip():
                extracted_text += f"\\n\\n--- Page {page_num + 1} ---\\n\\n"
                extracted_text += page_text
                
                # Create chunks for each page
                chunks.extend(self._create_chunks(page_text, page_num + 1, 'page'))
        
        doc.close()
        
        return {
            'success': True,
            'extracted_text': extracted_text,
            'chunks': chunks,
            'metadata': {
                'total_pages': len(doc),
                'format': 'PDF'
            }
        }
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX using python-docx"""
        doc = DocxDocument(file_path)
        extracted_text = ""
        chunks = []
        
        # Extract paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                extracted_text += paragraph.text + "\\n\\n"
                chunks.extend(self._create_chunks(paragraph.text, None, 'paragraph'))
        
        # Extract tables
        for table_num, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            if table_text.strip():
                extracted_text += f"\\n\\n--- Table {table_num + 1} ---\\n\\n"
                extracted_text += table_text + "\\n\\n"
                chunks.extend(self._create_chunks(table_text, None, 'table'))
        
        return {
            'success': True,
            'extracted_text': extracted_text,
            'chunks': chunks,
            'metadata': {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'format': 'DOCX'
            }
        }
    
    def _process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from XLSX using pandas and openpyxl"""
        workbook = load_workbook(file_path, data_only=True)
        extracted_text = ""
        chunks = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            extracted_text += f"\\n\\n--- Sheet: {sheet_name} ---\\n\\n"
            
            # Convert sheet to pandas DataFrame for easier processing
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Extract column headers
            headers = df.columns.tolist()
            header_text = "Columns: " + ", ".join(str(h) for h in headers)
            extracted_text += header_text + "\\n\\n"
            chunks.extend(self._create_chunks(header_text, None, 'headers'))
            
            # Extract data rows (limit to first 100 rows to avoid overwhelming)
            for idx, row in df.head(100).iterrows():
                row_text = " | ".join(str(val) for val in row.values if pd.notna(val))
                if row_text.strip():
                    extracted_text += row_text + "\\n"
                    chunks.extend(self._create_chunks(row_text, None, 'data_row'))
        
        workbook.close()
        
        return {
            'success': True,
            'extracted_text': extracted_text,
            'chunks': chunks,
            'metadata': {
                'total_sheets': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames,
                'format': 'XLSX'
            }
        }
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a DOCX table"""
        table_text = ""
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_text += row_text + "\\n"
        return table_text
    
    def _create_chunks(self, text: str, page_number: int = None, section_type: str = 'paragraph') -> List[Dict[str, Any]]:
        """Split text into smaller chunks for better processing"""
        # Clean and normalize text
        text = re.sub(r'\\s+', ' ', text).strip()
        
        if len(text) < 50:  # Skip very short chunks
            return []
        
        chunks = []
        
        # Split into sentences for better chunking
        sentences = re.split(r'[.!?]+', text)
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # If adding this sentence would make chunk too long, save current chunk
            if len(current_chunk) + len(sentence) > 500 and current_chunk:
                chunks.append({
                    'text': current_chunk.strip(),
                    'page_number': page_number,
                    'section_type': section_type,
                    'length': len(current_chunk.strip())
                })
                current_chunk = sentence
            else:
                current_chunk += (" " + sentence if current_chunk else sentence)
        
        # Add the last chunk if it exists
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'page_number': page_number,
                'section_type': section_type,
                'length': len(current_chunk.strip())
            })
        
        return chunks

class EmbeddingService:
    """Handles text embedding generation for semantic search"""
    
    def __init__(self):
        from sentence_transformers import SentenceTransformer
        # Use a lightweight but effective model
        self.model = SentenceTransformer('all-MiniLM-L6-v2')
    
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for a list of texts"""
        if not texts:
            return []
        
        embeddings = self.model.encode(texts, convert_to_tensor=False)
        return embeddings.tolist()
    
    def generate_single_embedding(self, text: str) -> List[float]:
        """Generate embedding for a single text"""
        embedding = self.model.encode([text], convert_to_tensor=False)
        return embedding[0].tolist()

class VectorStore:
    """Handles vector storage and similarity search using ChromaDB"""
    
    def __init__(self, persist_directory: str = "./chroma_db"):
        import chromadb
        from chromadb.config import Settings
        
        self.client = chromadb.PersistentClient(
            path=persist_directory,
            settings=Settings(anonymized_telemetry=False)
        )
        self.collection = self.client.get_or_create_collection(
            name="document_chunks",
            metadata={"hnsw:space": "cosine"}
        )
    
    def add_chunks(self, document_id: int, chunks: List[Dict[str, Any]], embeddings: List[List[float]]):
        """Add document chunks with their embeddings to the vector store"""
        if not chunks or not embeddings:
            return
        
        ids = [f"doc_{document_id}_chunk_{i}" for i in range(len(chunks))]
        documents = [chunk['text'] for chunk in chunks]
        metadatas = [
            {
                'document_id': document_id,
                'page_number': chunk.get('page_number'),
                'section_type': chunk.get('section_type', 'paragraph'),
                'length': chunk.get('length', len(chunk['text']))
            }
            for chunk in chunks
        ]
        
        self.collection.add(
            ids=ids,
            documents=documents,
            embeddings=embeddings,
            metadatas=metadatas
        )
    
    def search_similar(self, query_embedding: List[float], document_id: int = None, n_results: int = 5) -> Dict[str, Any]:
        """Search for similar chunks based on query embedding"""
        where_filter = {}
        if document_id:
            where_filter['document_id'] = document_id
        
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=n_results,
            where=where_filter if where_filter else None
        )
        
        return {
            'documents': results['documents'][0] if results['documents'] else [],
            'metadatas': results['metadatas'][0] if results['metadatas'] else [],
            'distances': results['distances'][0] if results['distances'] else []
        }
    
    def delete_document_chunks(self, document_id: int):
        """Delete all chunks for a specific document"""
        # Get all chunk IDs for this document
        results = self.collection.get(
            where={'document_id': document_id}
        )
        
        if results['ids']:
            self.collection.delete(ids=results['ids'])

